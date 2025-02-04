


from PyQt6 import QtCore, QtGui, QtWidgets
from PyQt6.QtWidgets import *
import psycopg2
from PyQt6.QtCore import QTimer
from PyQt6.QtGui import QFont
from docx import Document

class Ui_MainWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(1021, 806)
        self.centralwidget = QtWidgets.QWidget(parent=MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.widget = QtWidgets.QWidget(parent=self.centralwidget)
        self.widget.setGeometry(QtCore.QRect(10, 10, 1001, 781))
        self.widget.setObjectName("widget")
        self.horizontalLayout = QtWidgets.QHBoxLayout(self.widget)
        self.horizontalLayout.setContentsMargins(0, 0, 0, 0)
        self.horizontalLayout.setSpacing(19)
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.widget_2 = QtWidgets.QWidget(parent=self.widget)
        self.widget_2.setMinimumSize(QtCore.QSize(212, 0))
        self.widget_2.setStyleSheet("QWidget{\n"
"    \n"
"    background-color: rgb(0, 0, 59);\n"
"}\n"
"QPushButton{\n"
"    background-color: rgb(0, 0, 59);\n"
"    color: white;\n"
"    font: 11pt \"MS Shell Dlg 2\";\n"
"}\n"
"\n"
"")
        self.widget_2.setObjectName("widget_2")
        self.widget1 = QtWidgets.QWidget(parent=self.widget_2)
        self.widget1.setGeometry(QtCore.QRect(3, 30, 201, 784))
        self.widget1.setObjectName("widget1")
        self.verticalLayout = QtWidgets.QVBoxLayout(self.widget1)
        self.verticalLayout.setContentsMargins(0, 0, 0, 0)
        self.verticalLayout.setSpacing(0)
        self.verticalLayout.setObjectName("verticalLayout")
        self.homepushButton = QtWidgets.QPushButton(parent=self.widget1)
        self.homepushButton.setMinimumSize(QtCore.QSize(0, 48))
        self.homepushButton.setCheckable(True)
        self.homepushButton.setAutoExclusive(True)
        self.homepushButton.setObjectName("homepushButton")
        self.homepushButton.clicked.connect(lambda:self.stackedWidget.setCurrentIndex(0))
        self.verticalLayout.addWidget(self.homepushButton)
        self.bookAridepushbutton = QtWidgets.QPushButton(parent=self.widget1)
        self.bookAridepushbutton.setMinimumSize(QtCore.QSize(0, 47))
        self.bookAridepushbutton.setCheckable(True)
        self.bookAridepushbutton.setAutoExclusive(False)
        self.bookAridepushbutton.setObjectName("bookAridepushbutton")
        self.bookAridepushbutton.clicked.connect(lambda:self.stackedWidget.setCurrentIndex(1))
        self.verticalLayout.addWidget(self.bookAridepushbutton)
        self.myBookingpushbutton = QtWidgets.QPushButton(parent=self.widget1)
        self.myBookingpushbutton.setMinimumSize(QtCore.QSize(0, 47))
        self.myBookingpushbutton.setObjectName("myBookingpushbutton")
        self.myBookingpushbutton.clicked.connect(lambda:self.stackedWidget.setCurrentIndex(2))
        self.verticalLayout.addWidget(self.myBookingpushbutton)
        self.paymentspushbutton = QtWidgets.QPushButton(parent=self.widget1)
        self.paymentspushbutton.setMinimumSize(QtCore.QSize(0, 48))
        self.paymentspushbutton.setObjectName("paymentspushbutton")
        self.paymentspushbutton.clicked.connect(lambda:self.stackedWidget.setCurrentIndex(3))
        self.verticalLayout.addWidget(self.paymentspushbutton)
        self.profilepushbutton = QtWidgets.QPushButton(parent=self.widget1)
        self.profilepushbutton.setMinimumSize(QtCore.QSize(0, 48))
        self.profilepushbutton.setObjectName("profilepushbutton")
        self.profilepushbutton.clicked.connect(lambda:self.stackedWidget.setCurrentIndex(4))
        self.verticalLayout.addWidget(self.profilepushbutton)
        self.notificationpushbutton = QtWidgets.QPushButton(parent=self.widget1)
        self.notificationpushbutton.setMinimumSize(QtCore.QSize(0, 48))
        self.notificationpushbutton.setObjectName("notificationpushbutton")
        self.notificationpushbutton.clicked.connect(lambda:self.stackedWidget.setCurrentIndex(5))
        self.verticalLayout.addWidget(self.notificationpushbutton)
        self.helpsupprtPhbutton = QtWidgets.QPushButton(parent=self.widget1)
        self.helpsupprtPhbutton.setObjectName("helpsupprtPhbutton")
        self.helpsupprtPhbutton.clicked.connect(lambda:self.stackedWidget.setCurrentIndex(6))
        self.verticalLayout.addWidget(self.helpsupprtPhbutton)
        self.logutpushbutton = QtWidgets.QPushButton(parent=self.widget1)
        self.logutpushbutton.setObjectName("logutpushbutton")
        self.verticalLayout.addWidget(self.logutpushbutton)
        self.horizontalLayout.addWidget(self.widget_2)
        self.widget2 = QtWidgets.QWidget(parent=self.widget)
        self.widget2.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.widget2.setObjectName("widget2")
        self.gridLayout = QtWidgets.QGridLayout(self.widget2)
        self.gridLayout.setObjectName("gridLayout")
        self.stackedWidget = QtWidgets.QStackedWidget(parent=self.widget2)
        self.stackedWidget.setObjectName("stackedWidget")
        self.homepage = QtWidgets.QWidget()
        self.homepage.setObjectName("homepage")
        self.label = QtWidgets.QLabel(parent=self.homepage)
        self.label.setGeometry(QtCore.QRect(140, -20, 471, 71))
        self.label.setStyleSheet("\n"
"font: 15pt \"MS Shell Dlg 2\";")
        self.label.setObjectName("label")
        self.label_4 = QtWidgets.QLabel(parent=self.homepage)
        self.label_4.setGeometry(QtCore.QRect(190, 30, 301, 31))
        self.label_4.setStyleSheet("\n"
"font: 10pt \"MS Shell Dlg 2\";")
        self.label_4.setObjectName("label_4")
        self.lcdNumber = QtWidgets.QLCDNumber(parent=self.homepage)
        self.lcdNumber.setGeometry(QtCore.QRect(610, 650, 101, 81))
        self.lcdNumber.setObjectName("lcdNumber")
        self.label_5 = QtWidgets.QLabel(parent=self.homepage)
        self.label_5.setGeometry(QtCore.QRect(500, 620, 241, 31))
        self.label_5.setStyleSheet("\n"
"font: 10pt \"MS Shell Dlg 2\";")
        self.label_5.setObjectName("label_5")
        self.label_6 = QtWidgets.QLabel(parent=self.homepage)
        self.label_6.setGeometry(QtCore.QRect(10, 100, 301, 31))
        self.label_6.setStyleSheet("\n"
"font: 10pt \"MS Shell Dlg 2\";")
        self.label_6.setObjectName("label_6")
        self.hometableWidget = QtWidgets.QTableWidget(parent=self.homepage)
        self.hometableWidget.setGeometry(QtCore.QRect(0, 140, 741, 251))
        self.hometableWidget.setObjectName("hometableWidget")
        self.hometableWidget.setColumnCount(6)
        self.hometableWidget.setRowCount(0)
        item = QtWidgets.QTableWidgetItem()
        self.hometableWidget.setHorizontalHeaderItem(0, item)
        item = QtWidgets.QTableWidgetItem()
        self.hometableWidget.setHorizontalHeaderItem(1, item)
        item = QtWidgets.QTableWidgetItem()
        self.hometableWidget.setHorizontalHeaderItem(2, item)
        
        
   
        self.updatelcdDisplay()
        self.populatehometableWidget()
        self.createtablesifNotExists()
        self.timer=QTimer()
        self.timer.timeout.connect(self.populatehometableWidget)
        self.timer.timeout.connect(self.updatelcdDisplay)
        self.timer.start(10000)
        
        item = QtWidgets.QTableWidgetItem()
        self.hometableWidget.setHorizontalHeaderItem(3, item)
        item = QtWidgets.QTableWidgetItem()
        self.hometableWidget.setHorizontalHeaderItem(4, item)
        item = QtWidgets.QTableWidgetItem()
        self.hometableWidget.setHorizontalHeaderItem(5, item)
        self.label_7 = QtWidgets.QLabel(parent=self.homepage)
        self.label_7.setGeometry(QtCore.QRect(0, 400, 431, 31))
        self.label_7.setStyleSheet("\n"
"font: 10pt \"MS Shell Dlg 2\";")
        self.label_7.setObjectName("label_7")
        self.plainTextEdit = QtWidgets.QPlainTextEdit(parent=self.homepage)
        self.plainTextEdit.setGeometry(QtCore.QRect(0, 430, 741, 87))
        self.plainTextEdit.setObjectName("plainTextEdit")
        self.SUBMITcOMPLAINpushButton = QtWidgets.QPushButton(parent=self.homepage)
        self.SUBMITcOMPLAINpushButton.setGeometry(QtCore.QRect(0, 530, 93, 28))
        self.SUBMITcOMPLAINpushButton.setObjectName("SUBMITcOMPLAINpushButton")
        self.SUBMITcOMPLAINpushButton.clicked.connect(self.SUBMITcOMPLAINpushButtonClicked)
        self.label_8 = QtWidgets.QLabel(parent=self.homepage)
        self.label_8.setGeometry(QtCore.QRect(0, 730, 431, 31))
        self.label_8.setStyleSheet("\n"
"font: 10pt \"MS Shell Dlg 2\";")
        self.label_8.setObjectName("label_8")
        self.stackedWidget.addWidget(self.homepage)
        self.bookAridepage = QtWidgets.QWidget()
        self.bookAridepage.setObjectName("bookAridepage")
        self.label_2 = QtWidgets.QLabel(parent=self.bookAridepage)
        self.label_2.setGeometry(QtCore.QRect(170, 10, 411, 41))
        self.label_2.setStyleSheet("font: 12pt \"MS Shell Dlg 2\";")
        self.label_2.setObjectName("label_2")
        self.label_9 = QtWidgets.QLabel(parent=self.bookAridepage)
        self.label_9.setGeometry(QtCore.QRect(30, 110, 71, 41))
        self.label_9.setStyleSheet("font: 11pt \"MS Shell Dlg 2\";")
        self.label_9.setObjectName("label_9")
        self.FromcomboBox = QtWidgets.QComboBox(parent=self.bookAridepage)
        self.FromcomboBox.setGeometry(QtCore.QRect(110, 120, 291, 31))
        self.FromcomboBox.setObjectName("FromcomboBox")
        self.FromcomboBox.addItem("")
        self.FromcomboBox.addItem("")
        self.FromcomboBox.addItem("")
        self.FromcomboBox.addItem("")
        self.FromcomboBox.addItem("")
        self.FromcomboBox.addItem("")
        self.FromcomboBox.addItem("")
        self.FromcomboBox.addItem("")
        self.FromcomboBox.addItem("")
        self.label_10 = QtWidgets.QLabel(parent=self.bookAridepage)
        self.label_10.setGeometry(QtCore.QRect(30, 180, 71, 41))
        self.label_10.setStyleSheet("font: 11pt \"MS Shell Dlg 2\";")
        self.label_10.setObjectName("label_10")
        self.toComboBox = QtWidgets.QComboBox(parent=self.bookAridepage)
        self.toComboBox.setGeometry(QtCore.QRect(110, 190, 291, 31))
        self.toComboBox.setObjectName("toComboBox")
        self.toComboBox.addItem("")
        self.toComboBox.addItem("")
        self.toComboBox.addItem("")
        self.toComboBox.addItem("")
        self.toComboBox.addItem("")
        self.toComboBox.addItem("")
        self.toComboBox.addItem("")
        self.toComboBox.addItem("")
        self.toComboBox.addItem("")
        self.label_11 = QtWidgets.QLabel(parent=self.bookAridepage)
        self.label_11.setGeometry(QtCore.QRect(30, 270, 121, 41))
        self.label_11.setStyleSheet("font: 11pt \"MS Shell Dlg 2\";")
        self.label_11.setObjectName("label_11")
        self.TRAVELdAETdateEdit = QtWidgets.QDateEdit(parent=self.bookAridepage)
        self.TRAVELdAETdateEdit.setGeometry(QtCore.QRect(170, 280, 231, 22))
        self.TRAVELdAETdateEdit.setObjectName("TRAVELdAETdateEdit")
        self.label_12 = QtWidgets.QLabel(parent=self.bookAridepage)
        self.label_12.setGeometry(QtCore.QRect(30, 340, 161, 41))
        self.label_12.setStyleSheet("font: 11pt \"MS Shell Dlg 2\";")
        self.label_12.setObjectName("label_12")
        self.NumberOfSeatsCombobox = QtWidgets.QComboBox(parent=self.bookAridepage)
        self.NumberOfSeatsCombobox.setGeometry(QtCore.QRect(200, 340, 291, 31))
        self.NumberOfSeatsCombobox.setObjectName("NumberOfSeatsCombobox")
        self.NumberOfSeatsCombobox.addItem("")
        self.NumberOfSeatsCombobox.addItem("")
        self.NumberOfSeatsCombobox.addItem("")
        self.NumberOfSeatsCombobox.addItem("")
        self.NumberOfSeatsCombobox.addItem("")
        self.NumberOfSeatsCombobox.addItem("")
        self.label_13 = QtWidgets.QLabel(parent=self.bookAridepage)
        self.label_13.setGeometry(QtCore.QRect(30, 410, 131, 41))
        self.label_13.setStyleSheet("font: 11pt \"MS Shell Dlg 2\";")
        self.label_13.setObjectName("label_13")
        self.BUScATEGORYcomboBox = QtWidgets.QComboBox(parent=self.bookAridepage)
        self.BUScATEGORYcomboBox.setGeometry(QtCore.QRect(200, 420, 291, 31))
        self.BUScATEGORYcomboBox.setObjectName("BUScATEGORYcomboBox")
        self.BUScATEGORYcomboBox.addItem("")
        self.BUScATEGORYcomboBox.addItem("")
        self.BUScATEGORYcomboBox.addItem("")
        self.label_14 = QtWidgets.QLabel(parent=self.bookAridepage)
        self.label_14.setGeometry(QtCore.QRect(30, 480, 121, 41))
        self.label_14.setStyleSheet("font: 11pt \"MS Shell Dlg 2\";")
        self.label_14.setObjectName("label_14")
        self.NAMElineEdit = QtWidgets.QLineEdit(parent=self.bookAridepage)
        self.NAMElineEdit.setGeometry(QtCore.QRect(200, 480, 291, 31))
        self.NAMElineEdit.setObjectName("NAMElineEdit")
        self.label_15 = QtWidgets.QLabel(parent=self.bookAridepage)
        self.label_15.setGeometry(QtCore.QRect(30, 540, 141, 41))
        self.label_15.setStyleSheet("font: 11pt \"MS Shell Dlg 2\";")
        self.label_15.setObjectName("label_15")
        self.PHONEnumberlineedit = QtWidgets.QLineEdit(parent=self.bookAridepage)
        self.PHONEnumberlineedit.setGeometry(QtCore.QRect(200, 540, 291, 31))
        self.PHONEnumberlineedit.setObjectName("PHONEnumberlineedit")
        self.BOOKpushButton = QtWidgets.QPushButton(parent=self.bookAridepage)
        self.BOOKpushButton.setGeometry(QtCore.QRect(200, 620, 93, 28))
        self.BOOKpushButton.setObjectName("BOOKpushButton")
        self.ShwbookingsummaryButton = QtWidgets.QPushButton(parent=self.bookAridepage)
        self.ShwbookingsummaryButton.setGeometry(QtCore.QRect(330, 620, 161, 28))
        self.ShwbookingsummaryButton.setObjectName("ShwbookingsummaryButton")
        self.stackedWidget.addWidget(self.bookAridepage)
        self.bookingsPage = QtWidgets.QWidget()
        self.bookingsPage.setObjectName("bookingsPage")
        self.label_3 = QtWidgets.QLabel(parent=self.bookingsPage)
        self.label_3.setGeometry(QtCore.QRect(150, -10, 331, 71))
        self.label_3.setStyleSheet("font: 13pt \"MS Shell Dlg 2\";")
        self.label_3.setObjectName("label_3")
        self.myBookingstableWidget_2 = QtWidgets.QTableWidget(parent=self.bookingsPage)
        self.myBookingstableWidget_2.setGeometry(QtCore.QRect(0, 120, 751, 211))
        self.myBookingstableWidget_2.setObjectName("myBookingstableWidget_2")
        self.myBookingstableWidget_2.setColumnCount(6)
        self.myBookingstableWidget_2.setRowCount(0)
        item = QtWidgets.QTableWidgetItem()
        self.myBookingstableWidget_2.setHorizontalHeaderItem(0, item)
        item = QtWidgets.QTableWidgetItem()
        self.myBookingstableWidget_2.setHorizontalHeaderItem(1, item)
        item = QtWidgets.QTableWidgetItem()
        self.myBookingstableWidget_2.setHorizontalHeaderItem(2, item)
        item = QtWidgets.QTableWidgetItem()
        self.myBookingstableWidget_2.setHorizontalHeaderItem(3, item)
        item = QtWidgets.QTableWidgetItem()
        self.myBookingstableWidget_2.setHorizontalHeaderItem(4, item)
        item = QtWidgets.QTableWidgetItem()
        self.myBookingstableWidget_2.setHorizontalHeaderItem(5, item)
        self.cancelBokingcommandLinkButton = QtWidgets.QCommandLinkButton(parent=self.bookingsPage)
        self.cancelBokingcommandLinkButton.setGeometry(QtCore.QRect(0, 370, 431, 48))
        self.cancelBokingcommandLinkButton.setObjectName("cancelBokingcommandLinkButton")
        self.stackedWidget.addWidget(self.bookingsPage)
        self.paymentPage = QtWidgets.QWidget()
        self.paymentPage.setObjectName("paymentPage")
        self.label_16 = QtWidgets.QLabel(parent=self.paymentPage)
        self.label_16.setGeometry(QtCore.QRect(0, 0, 401, 31))
        self.label_16.setStyleSheet("font: 12pt \"MS Shell Dlg 2\";")
        self.label_16.setObjectName("label_16")
        self.label_17 = QtWidgets.QLabel(parent=self.paymentPage)
        self.label_17.setGeometry(QtCore.QRect(0, 70, 71, 41))
        self.label_17.setStyleSheet("font: 12pt \"MS Shell Dlg 2\";")
        self.label_17.setObjectName("label_17")
        self.fromPaymentsSeclineEdit = QtWidgets.QLineEdit(parent=self.paymentPage)
        self.fromPaymentsSeclineEdit.setGeometry(QtCore.QRect(90, 71, 271, 31))
        self.fromPaymentsSeclineEdit.setObjectName("fromPaymentsSeclineEdit")
        self.label_18 = QtWidgets.QLabel(parent=self.paymentPage)
        self.label_18.setGeometry(QtCore.QRect(0, 140, 71, 41))
        self.label_18.setStyleSheet("font: 12pt \"MS Shell Dlg 2\";")
        self.label_18.setObjectName("label_18")
        self.toPaymentsSeclineEdit = QtWidgets.QLineEdit(parent=self.paymentPage)
        self.toPaymentsSeclineEdit.setGeometry(QtCore.QRect(90, 140, 271, 31))
        self.toPaymentsSeclineEdit.setObjectName("toPaymentsSeclineEdit")
        self.label_19 = QtWidgets.QLabel(parent=self.paymentPage)
        self.label_19.setGeometry(QtCore.QRect(0, 220, 71, 41))
        self.label_19.setStyleSheet("font: 12pt \"MS Shell Dlg 2\";")
        self.label_19.setObjectName("label_19")
        self.seatsPaymentsSeclineEdit = QtWidgets.QLineEdit(parent=self.paymentPage)
        self.seatsPaymentsSeclineEdit.setGeometry(QtCore.QRect(90, 220, 271, 31))
        self.seatsPaymentsSeclineEdit.setObjectName("seatsPaymentsSeclineEdit")
        self.label_20 = QtWidgets.QLabel(parent=self.paymentPage)
        self.label_20.setGeometry(QtCore.QRect(0, 300, 181, 41))
        self.label_20.setStyleSheet("font: 12pt \"MS Shell Dlg 2\";")
        self.label_20.setObjectName("label_20")
        self.paymentcomboBox = QtWidgets.QComboBox(parent=self.paymentPage)
        self.paymentcomboBox.setGeometry(QtCore.QRect(190, 310, 211, 22))
        self.paymentcomboBox.setObjectName("paymentcomboBox")
        self.paymentcomboBox.addItem("")
        self.paymentcomboBox.addItem("")
        self.paymentcomboBox.addItem("")
        self.paymentcomboBox.addItem("")
        self.label_21 = QtWidgets.QLabel(parent=self.paymentPage)
        self.label_21.setGeometry(QtCore.QRect(0, 360, 121, 51))
        self.label_21.setStyleSheet("font: 12pt \"MS Shell Dlg 2\";")
        self.label_21.setObjectName("label_21")
        self.totalFareLinedit = QtWidgets.QLineEdit(parent=self.paymentPage)
        self.totalFareLinedit.setGeometry(QtCore.QRect(150, 370, 271, 31))
        self.totalFareLinedit.setObjectName("totalFareLinedit")
        self.paypushButton = QtWidgets.QPushButton(parent=self.paymentPage)
        self.paypushButton.setGeometry(QtCore.QRect(160, 440, 141, 28))
        self.paypushButton.setObjectName("paypushButton")
        self.stackedWidget.addWidget(self.paymentPage)
        self.profilepage = QtWidgets.QWidget()
        self.profilepage.setObjectName("profilepage")
        self.label_22 = QtWidgets.QLabel(parent=self.profilepage)
        self.label_22.setGeometry(QtCore.QRect(10, 20, 251, 16))
        self.label_22.setStyleSheet("font: 12pt \"MS Shell Dlg 2\";")
        self.label_22.setObjectName("label_22")
        self.label_23 = QtWidgets.QLabel(parent=self.profilepage)
        self.label_23.setGeometry(QtCore.QRect(0, 100, 101, 16))
        self.label_23.setStyleSheet("font: 12pt \"MS Shell Dlg 2\";")
        self.label_23.setObjectName("label_23")
        self.namelineEdit_profilepage = QtWidgets.QLineEdit(parent=self.profilepage)
        self.namelineEdit_profilepage.setGeometry(QtCore.QRect(120, 91, 341, 31))
        self.namelineEdit_profilepage.setObjectName("namelineEdit_profilepage")
        self.label_24 = QtWidgets.QLabel(parent=self.profilepage)
        self.label_24.setGeometry(QtCore.QRect(0, 170, 131, 16))
        self.label_24.setStyleSheet("font: 12pt \"MS Shell Dlg 2\";")
        self.label_24.setObjectName("label_24")
        self.phonenumberedit_profilepage = QtWidgets.QLineEdit(parent=self.profilepage)
        self.phonenumberedit_profilepage.setGeometry(QtCore.QRect(140, 160, 341, 31))
        self.phonenumberedit_profilepage.setObjectName("phonenumberedit_profilepage")
        self.savepushButton_profilepage = QtWidgets.QPushButton(parent=self.profilepage)
        
        self.savepushButton_profilepage.clicked.connect(self.savepushButton_profilepageClicked)
        
        self.savepushButton_profilepage.setGeometry(QtCore.QRect(150, 240, 93, 28))
        self.savepushButton_profilepage.setObjectName("savepushButton_profilepage")
        self.stackedWidget.addWidget(self.profilepage)
        self.notificationPage = QtWidgets.QWidget()
        self.notificationPage.setObjectName("notificationPage")
        self.notificationplainTextEdit_notificationpage = QtWidgets.QPlainTextEdit(parent=self.notificationPage)
        self.notificationplainTextEdit_notificationpage.setGeometry(QtCore.QRect(0, 50, 751, 521))
        self.notificationplainTextEdit_notificationpage.setObjectName("notificationplainTextEdit_notificationpage")
        self.label_25 = QtWidgets.QLabel(parent=self.notificationPage)
        
        self.setTexttonotificationplainTextEdit_notificationpage()
        
        self.label_25.setGeometry(QtCore.QRect(40, 10, 251, 16))
        self.label_25.setStyleSheet("font: 11pt \"MS Shell Dlg 2\";")
        self.label_25.setObjectName("label_25")
        self.stackedWidget.addWidget(self.notificationPage)
        self.helpandSupportPage = QtWidgets.QWidget()
        self.helpandSupportPage.setObjectName("helpandSupportPage")
        self.label_26 = QtWidgets.QLabel(parent=self.helpandSupportPage)
        self.label_26.setGeometry(QtCore.QRect(20, 0, 101, 16))
        self.label_26.setStyleSheet("font: 11pt \"MS Shell Dlg 2\";")
        self.label_26.setObjectName("label_26")
        self.textEdit = QtWidgets.QTextEdit(parent=self.helpandSupportPage)
        self.textEdit.setGeometry(QtCore.QRect(0, 50, 751, 301))
        self.textEdit.setObjectName("textEdit")
        self.helponHowToUsesystemcommandLinkButton = QtWidgets.QCommandLinkButton(parent=self.helpandSupportPage)
        self.helponHowToUsesystemcommandLinkButton.setGeometry(QtCore.QRect(10, 380, 321, 48))
        self.helponHowToUsesystemcommandLinkButton.setObjectName("helponHowToUsesystemcommandLinkButton")
        self.reportanIssuecommandLinkButton = QtWidgets.QCommandLinkButton(parent=self.helpandSupportPage)
        self.reportanIssuecommandLinkButton.setGeometry(QtCore.QRect(10, 440, 222, 48))
        self.reportanIssuecommandLinkButton.setObjectName("reportanIssuecommandLinkButton")
        self.stackedWidget.addWidget(self.helpandSupportPage)
        self.gridLayout.addWidget(self.stackedWidget, 0, 0, 1, 1)
        self.horizontalLayout.addWidget(self.widget2)
        MainWindow.setCentralWidget(self.centralwidget)
        self.statusbar = QtWidgets.QStatusBar(parent=MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)
        
        self.logutpushbutton.clicked.connect(self.logouttoLoginWindow)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)


    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "Traveller Dashboard-QuickRideCompany"))
        self.homepushButton.setText(_translate("MainWindow", "Home"))
        self.bookAridepushbutton.setText(_translate("MainWindow", "Book a Ride"))
        self.myBookingpushbutton.setText(_translate("MainWindow", "My Bookigs"))
        self.paymentspushbutton.setText(_translate("MainWindow", "Payments"))
        self.profilepushbutton.setText(_translate("MainWindow", "Profile"))
        self.notificationpushbutton.setText(_translate("MainWindow", "Notifications"))
        self.helpsupprtPhbutton.setText(_translate("MainWindow", "Help Support"))
        self.logutpushbutton.setText(_translate("MainWindow", "Logout"))
        self.label.setText(_translate("MainWindow", "WELCOME TO QUICKRIDECOMPANY"))
        self.label_4.setText(_translate("MainWindow", "Reach Your Destination On Time"))
        self.label_5.setText(_translate("MainWindow", "AVAILABLE BUSES FOR A RIDE"))
        self.label_6.setText(_translate("MainWindow", "TRIPS AVAILABLE TODAY"))
        item = self.hometableWidget.horizontalHeaderItem(0)
        item.setText(_translate("MainWindow", "BUS REG"))
        item = self.hometableWidget.horizontalHeaderItem(1)
        item.setText(_translate("MainWindow", "DEPARTURE TIME"))
        item = self.hometableWidget.horizontalHeaderItem(2)
        item.setText(_translate("MainWindow", "DEPARTURE FROM"))
        item = self.hometableWidget.horizontalHeaderItem(3)
        item.setText(_translate("MainWindow", "DESTINATION"))
        item = self.hometableWidget.horizontalHeaderItem(4)
        item.setText(_translate("MainWindow", "AVAILABLE SEATS"))
        item = self.hometableWidget.horizontalHeaderItem(5)
        item.setText(_translate("MainWindow", "COST"))
        self.label_7.setText(_translate("MainWindow", "YOU HAVE A COMPLAIN/COMMENT ON OUR SERVICES?"))
        self.SUBMITcOMPLAINpushButton.setText(_translate("MainWindow", "SUBMIT"))
        self.label_8.setText(_translate("MainWindow", "Connecting You to Your Journey, Effortlessly."))
        self.label_2.setText(_translate("MainWindow", "Book Your Ride: Fast, Easy, and Convenient"))
        self.label_9.setText(_translate("MainWindow", "FROM"))
        self.FromcomboBox.setItemText(0, _translate("MainWindow", "MOMBASA"))
        self.FromcomboBox.setItemText(1, _translate("MainWindow", "NAIROBI"))
        self.FromcomboBox.setItemText(2, _translate("MainWindow", "NAKURU"))
        self.FromcomboBox.setItemText(3, _translate("MainWindow", "KISII"))
        self.FromcomboBox.setItemText(4, _translate("MainWindow", "KAKAMEGA"))
        self.FromcomboBox.setItemText(5, _translate("MainWindow", "KISUMU"))
        self.FromcomboBox.setItemText(6, _translate("MainWindow", "NAIVASHA"))
        self.FromcomboBox.setItemText(7, _translate("MainWindow", "BUNGOMA"))
        self.FromcomboBox.setItemText(8, _translate("MainWindow", "LAMU"))
        self.label_10.setText(_translate("MainWindow", "TO"))
        self.toComboBox.setItemText(0, _translate("MainWindow", "MOMBASA"))
        self.toComboBox.setItemText(1, _translate("MainWindow", "NAIROBI"))
        self.toComboBox.setItemText(2, _translate("MainWindow", "KISII"))
        self.toComboBox.setItemText(3, _translate("MainWindow", "KISUMU"))
        self.toComboBox.setItemText(4, _translate("MainWindow", "NAKURU"))
        self.toComboBox.setItemText(5, _translate("MainWindow", "KAKAMEGA"))
        self.toComboBox.setItemText(6, _translate("MainWindow", "NAIVASHA"))
        self.toComboBox.setItemText(7, _translate("MainWindow", "BUNGOMA"))
        self.toComboBox.setItemText(8, _translate("MainWindow", "LAMU"))
        self.label_11.setText(_translate("MainWindow", "TRAVEL DATE"))
        self.label_12.setText(_translate("MainWindow", "NUMBER OF SEATS"))
        self.NumberOfSeatsCombobox.setItemText(0, _translate("MainWindow", "1"))
        self.NumberOfSeatsCombobox.setItemText(1, _translate("MainWindow", "2"))
        self.NumberOfSeatsCombobox.setItemText(2, _translate("MainWindow", "3"))
        self.NumberOfSeatsCombobox.setItemText(3, _translate("MainWindow", "4"))
        self.NumberOfSeatsCombobox.setItemText(4, _translate("MainWindow", "5"))
        self.NumberOfSeatsCombobox.setItemText(5, _translate("MainWindow", "Full Bus"))
        self.label_13.setText(_translate("MainWindow", "BUS CATEGORY"))
        self.BUScATEGORYcomboBox.setItemText(0, _translate("MainWindow", "STANDARD"))
        self.BUScATEGORYcomboBox.setItemText(1, _translate("MainWindow", "EXECUTIVE"))
        self.BUScATEGORYcomboBox.setItemText(2, _translate("MainWindow", "VIP"))
        self.label_14.setText(_translate("MainWindow", "NAME"))
        self.label_15.setText(_translate("MainWindow", "PHONE NUMBER"))
        self.BOOKpushButton.setText(_translate("MainWindow", "BOOK"))
        
        self.BOOKpushButton.clicked.connect(self.BOOKpushButtonClicked)
        
        self.ShwbookingsummaryButton.setText(_translate("MainWindow", "SHOW BOOKING SUMMARY"))
        self.ShwbookingsummaryButton.clicked.connect(self.ShwbookingsummaryButtonClicked)
        self.label_3.setText(_translate("MainWindow", "YOUR PREVIOUS BOOKINGS"))
        item = self.myBookingstableWidget_2.horizontalHeaderItem(0)
        item.setText(_translate("MainWindow", "FROM"))
        item = self.myBookingstableWidget_2.horizontalHeaderItem(1)
        item.setText(_translate("MainWindow", "TO"))
        item = self.myBookingstableWidget_2.horizontalHeaderItem(2)
        item.setText(_translate("MainWindow", "SEATS"))
        item = self.myBookingstableWidget_2.horizontalHeaderItem(3)
        item.setText(_translate("MainWindow", "COST"))
        item = self.myBookingstableWidget_2.horizontalHeaderItem(4)
        item.setText(_translate("MainWindow", "DEPARTURE TIME"))
        item = self.myBookingstableWidget_2.horizontalHeaderItem(5)
        item.setText(_translate("MainWindow", "BUS"))
        self.cancelBokingcommandLinkButton.setText(_translate("MainWindow", "CLICK HERE IF YOU WANT TO CANCEL THE BOOKING"))
        self.label_16.setText(_translate("MainWindow", "Make Payments Of The Booking You made"))
        self.label_17.setText(_translate("MainWindow", "FROM"))
        self.label_18.setText(_translate("MainWindow", "TO"))
        self.label_19.setText(_translate("MainWindow", "SEATS"))
        self.label_20.setText(_translate("MainWindow", "PAYMENT METHOD"))
        self.paymentcomboBox.setItemText(0, _translate("MainWindow", "M-PESA"))
        self.paymentcomboBox.setItemText(1, _translate("MainWindow", "AIRTEL MONEY"))
        self.paymentcomboBox.setItemText(2, _translate("MainWindow", "KCB"))
        self.paymentcomboBox.setItemText(3, _translate("MainWindow", "EQUITY"))
        self.label_21.setText(_translate("MainWindow", "TOTAL FARE"))
        self.paypushButton.setText(_translate("MainWindow", "PAY"))
        self.label_22.setText(_translate("MainWindow", "EDIT YOUR PROFILE"))
        self.label_23.setText(_translate("MainWindow", "NAME"))
        self.label_24.setText(_translate("MainWindow", "PHONE NO."))
        self.savepushButton_profilepage.setText(_translate("MainWindow", "SAVE"))
        self.label_25.setText(_translate("MainWindow", "TODAY ANOUNCEMENTS"))
        self.label_26.setText(_translate("MainWindow", "CREDITS"))
        self.textEdit.setHtml(_translate("MainWindow", "<!DOCTYPE HTML PUBLIC \"-//W3C//DTD HTML 4.0//EN\" \"http://www.w3.org/TR/REC-html40/strict.dtd\">\n"
"<html><head><meta name=\"qrichtext\" content=\"1\" /><style type=\"text/css\">\n"
"p, li { white-space: pre-wrap; }\n"
"</style></head><body style=\" font-family:\'MS Shell Dlg 2\'; font-size:7.8pt; font-weight:400; font-style:normal;\">\n"
"<p align=\"center\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:11pt; font-weight:600; font-style:italic;\">developed by ;</span></p>\n"
"<p align=\"center\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:11pt; font-weight:600; font-style:italic;\">Senior Developer- Justin Ratemo</span></p>\n"
"<p align=\"center\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:11pt; font-weight:600; font-style:italic;\">GUI developer- Amos Kilonzo</span></p>\n"
"<p align=\"center\" style=\"-qt-paragraph-type:empty; margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px; font-size:11pt; font-weight:600; font-style:italic;\"><br /></p>\n"
"<p align=\"center\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:11pt; font-weight:600; font-style:italic;\">Languages</span></p>\n"
"<p align=\"center\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:11pt; font-weight:600; font-style:italic;\">Python and Pyqt6 Framework</span></p>\n"
"<p align=\"center\" style=\"-qt-paragraph-type:empty; margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px; font-size:11pt; font-weight:600; font-style:italic;\"><br /></p>\n"
"<p align=\"center\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:11pt; font-weight:600; font-style:italic;\">Database used</span></p>\n"
"<p align=\"center\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:11pt; font-weight:600; font-style:italic;\"> PostgreSQL</span></p>\n"
"<p align=\"center\" style=\"-qt-paragraph-type:empty; margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px; font-size:11pt; font-weight:600; font-style:italic;\"><br /></p>\n"
"<p align=\"center\" style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\"><span style=\" font-size:11pt; font-weight:600; font-style:italic;\">PHONE NO. 0793031269</span></p>\n"
"<p align=\"center\" style=\"-qt-paragraph-type:empty; margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px; font-size:11pt; font-weight:600; font-style:italic;\"><br /></p></body></html>"))
        self.helponHowToUsesystemcommandLinkButton.setText(_translate("MainWindow", "Help on How to use system?"))
        self.reportanIssuecommandLinkButton.setText(_translate("MainWindow", "REPORT AN ISSUE"))
        """
        self.FromcomboBox.currentTextChanged.connect(self.BOOKpushButtonClicked)
        self.toComboBox.currentTextChanged.connect(self.BOOKpushButtonClicked)
        self.NumberOfSeatsCombobox.currentTextChanged.connect(self.BOOKpushButtonClicked)
        self.BUScATEGORYcomboBox.currentTextChanged.connect(self.BOOKpushButtonClicked)"""
        
    def SUBMITcOMPLAINpushButtonClicked(self):
        complain=self.plainTextEdit.toPlainText().strip()
        connection=self.connectDB()
        
        if connection:
                cursor=connection.cursor()
                cursor.execute("CREATE TABLE IF NOT EXISTS complain(complain VARCHAR(1000))")
                cursor.execute("INSERT INTO complain (complain) VALUES(%s)",(complain,))
                connection.commit()
                QMessageBox.information(None,"Success","Submitted Successsfully\nThanks For Your Feedback!")
                self.plainTextEdit.clear()
                connection.close()
                cursor.close()
        
        
     #db connection
    def connectDB(self):
            try:
                    connection=psycopg2.connect(
                            host="localhost",
                            user="postgres",
                            password="3062",
                            database="quickrideCompany"
                            )
                    
                    return connection
            except Exception as e:
                    QMessageBox.warning(None,"ERROR","ERROR OCCURED: ")
                    
    def createtablesifNotExists(self):
            connection=self.connectDB()
            if connection:
                    cursor=connection.cursor()
                    cursor.execute("""
                                   CREATE TABLE IF NOT EXISTS tripavailabletoday(
                                           busreg VARCHAR(20),
                                           departuretime VARCHAR(20),
                                           arrivaltime VARCHAR(20),
                                           destination VARCHAR(20),
                                           availableseats VARCHAR(20),
                                           cost INTEGER
                                   )
                                   
                                   """)
                    connection.commit()
                    connection.close()
                    cursor.close()



    def populatehometableWidget(self):
        connection=self.connectDB()
        if connection:
                cursor=connection.cursor()
                cursor.execute("SELECT busreg,departuretime,arrivaltime,destination,availableseats,cost FROM tripavailabletoday")
                trips=cursor.fetchall()
                cursor.close()
                connection.close()
                self.hometableWidget.setRowCount(len(trips))
                for row,rowdata in enumerate(trips):
                        for column,celldata in enumerate(rowdata):
                                self.hometableWidget.setItem(row,column,QTableWidgetItem(str(celldata)))
                                
                                
    def updatelcdDisplay(self):
        connection=self.connectDB()
        
        if connection:
                cursor=connection.cursor()
                cursor.execute("SELECT COUNT(*) FROM tripavailabletoday")
                count=cursor.fetchone()[0]
                self.lcdNumber.display(count)
                cursor.close()
                connection.close()
                
    def logouttoLoginWindow(self):
            from loginUi import Ui_loginDialog
            self.loginwindow = QtWidgets.QMainWindow()
            self.ui = Ui_loginDialog()
            self.ui.setupUi(self.loginwindow)
            self.loginwindow.setFixedSize(369, 308)
            self.loginwindow.show()
            QtWidgets.QApplication.instance().activeWindow().close()
            QMessageBox.information(None,"Logout Successfully","LOGGED OUT SUCCESSFULLY!")
            
    def setTexttonotificationplainTextEdit_notificationpage(self):
           connection=self.connectDB()
           if connection:
                   cursor=connection.cursor()
                   cursor.execute("SELECT complain FROM complain")
                   complains=cursor.fetchall()
                   self.notificationplainTextEdit_notificationpage.setReadOnly(True)
                   font=QFont("Arial",13)
                   formcomplain="\n".join(row[0] for row in complains)
                   self.notificationplainTextEdit_notificationpage.setFont(font)
                   self.notificationplainTextEdit_notificationpage.setPlainText("""Be Updated!\n\nFree water would be provided to all classes as from monday!\nDont accept any food from strangers!\n\n\nCOMPLAIN TO BE ADDRESSSED AS SUBMITTED BY CUSTOMERS\n""")
                  
                   self.notificationplainTextEdit_notificationpage.appendPlainText(formcomplain)
                   cursor.close()
                   connection.close()
    
    def savepushButton_profilepageClicked(self):
            name=self.namelineEdit_profilepage.text().strip()
            phoneNumber=self.phonenumberedit_profilepage.text().strip()
            connection=self.connectDB()
            if connection:
                    cursor=connection.cursor()
                    cursor.execute("CREATE TABLE IF NOT EXISTS profile (name VARCHAR(20),phoneno VARCHAR(20))")
                    cursor.execute("INSERT INTO profile (name,phoneno) VALUES(%s,%s)",(name,phoneNumber))
                    connection.commit()
                    QMessageBox.information(None,"Success","Profile Info Saved Successfully!")
                    """self.namelineEdit_profilepage.setReadOnly(True)
                    self.phonenumberedit_profilepage.setReadOnly(True)
                    self.savepushButton_profilepage.setCheckable(True)"""
                    cursor.close()
                    connection.close()
    
    def BOOKpushButtonClicked(self):
            From=self.FromcomboBox.currentText().lower()
            To=self.toComboBox.currentText().lower()
            numberofSeats=int(self.NumberOfSeatsCombobox.currentText())
            buscategory=self.BUScATEGORYcomboBox.currentText().lower()
            traveldate=self.TRAVELdAETdateEdit.text()
            name=self.NAMElineEdit.text()
            phoneno=self.PHONEnumberlineedit.text()
            booking_details=(
                    f"From: {From}\n"
                    f"To: {To}\n"
                    f"Number Of Seats: {numberofSeats}\n"
                    f"Bus Category: {buscategory}\n"
                    f"Travel Date: {traveldate}\n"
                    f"Name: {name}\n"
                    f"Phone No.: {phoneno}"
            )
            
            if not name or not phoneno:
                    QMessageBox.warning(None,"Error","All Fields are required !")
                    return
            else:
                connection=self.connectDB()
                if connection:
                    cursor=connection.cursor()
                    cursor.execute("""
                                   CREATE TABLE IF NOT EXISTS bookaridetable(
                                           fromplace VARCHAR(20),
                                          destination VARCHAR(20),
                                           traveldate VARCHAR(20),
                                           noofseats INTEGER,
                                           buscategory VARCHAR(20),
                                           name VARCHAR(20),
                                           phoneno VARCHAR(20)
                                           
                                   )
                                   """)
                    cursor.execute("""
                                   INSERT INTO bookaridetable (
                                           fromplace,destination,traveldate,noofseats,buscategory,name,phoneno
                                   )
                                   VALUES(%s,%s,%s,%s,%s,%s,%s)
                                   """,(From,To,traveldate,numberofSeats,buscategory,name,phoneno))
                    
                    connection.commit()
                    QMessageBox.information(None,"SUCCESS","BOOKED SUCCESSFUL!")
                    dialog=QDialog()
                    dialog.setWindowTitle("Booking Confirmation")
                    dialog.setGeometry(300,200,400,300)
                    layout=QVBoxLayout()
                    label=QLabel("YOUR BOOKING DETAILS")
                    layout.addWidget(label)
                    textEdit=QTextEdit()
                    textEdit.setReadOnly(True)
                    layout.addWidget(textEdit)
                    dialog.setLayout(layout)
                    textEdit.setText(booking_details)
                    dialog.exec()
                    #options=QFileDialog.options(0)
                    filePath, _ = QFileDialog.getSaveFileName( None,"Save Booking Info AS docx", "", "Word Documents (*.docx)")

                    if filePath:
                           progress = QProgressDialog("Saving booking details...", None, 0, 100)
                           progress.setWindowTitle("Saving")
                           progress.setCancelButton(None)
                           progress.setMinimumDuration(500)  
                           progress.show()
                           import time

    
                           for i in range(1, 101, 20):  
                                time.sleep(0.2)
                                progress.setValue(i)
                               
                           doc=Document()
                           doc.add_heading("Booking Confirmation",0)
                           doc.add_paragraph(booking_details)
                           doc.save(filePath)
                           progress.setValue(100)
                           import os
                           fileName = os.path.basename(filePath) 
                           QMessageBox.information(None,"File Saved", f"Booking info saved as '{fileName}' at:\n{filePath}")
                    else:
                        QMessageBox.warning(self, "Save Cancelled", "No file was saved.")
   

                    self.NAMElineEdit.clear()
                    self.PHONEnumberlineedit.clear()
                    cursor.close()
                    connection.close()
    def ShwbookingsummaryButtonClicked(self):
            From=self.FromcomboBox.currentText().lower()
            To=self.toComboBox.currentText().lower()
            numberofSeats=int(self.NumberOfSeatsCombobox.currentText())
            buscategory=self.BUScATEGORYcomboBox.currentText().lower()
            traveldate=self.TRAVELdAETdateEdit.text()
            name=self.NAMElineEdit.text()
            phoneno=self.PHONEnumberlineedit.text()
            booking_details=(
                    f"From: {From}\n"
                    f"To: {To}\n"
                    f"Number Of Seats: {numberofSeats}\n"
                    f"Bus Category: {buscategory}\n"
                    f"Travel Date: {traveldate}\n"
                    f"Name: {name}\n"
                    f"Phone No.: {phoneno}"
            )
            dialog=QDialog()
            dialog.setWindowTitle("Booking Summary")
            dialog.setGeometry(300,200,400,300)
            layout=QVBoxLayout()
            label=QLabel("YOUR BOOKING DETAILS")
            layout.addWidget(label)
            textEdit=QTextEdit()
            textEdit.setReadOnly(True)
            layout.addWidget(textEdit)
            dialog.setLayout(layout)
            textEdit.setText(booking_details)
            dialog.exec()
                    
            




if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_MainWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec())
